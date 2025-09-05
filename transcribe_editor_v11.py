# transcribe_editor_v11.py
# -------------------------------------------------------------
# 機能:
# 1) 音声/動画アップロード → OpenAIで文字起こし（50MBまで直接可、超過は自動分割）
# 2) 出力選択: 逐語(タイムスタンプ) / 直訳（日本語化のみ）/ 議事録 / 要旨 / 記事 / ガイドライン解説
# 3) 目的選択: 学会発表 / ガイドライン解説 / ディスカッション（LLM整形に反映）
# 4) 動画オプション: スライドOCR(キーフレーム抽出 + OCR) 併用の可否（依存が無ければ自動でスキップ）
# 5) TXT/DOCXでダウンロード可能
# 6) サイドバーで毎回「共通パスワード」と「OpenAI APIキー」を入力（Secrets不要）
# -------------------------------------------------------------

import os
import io
import time
import glob
import shutil
import subprocess
import mimetypes
import json
import traceback
from datetime import timedelta
import re
from typing import List, Tuple, Dict, Any

import math
from pathlib import Path

import streamlit as st
from pydub import AudioSegment
from pydub.utils import which
from docx import Document
from docx.shared import Pt

# （LLMの整形に chat.completions を使うための互換ハンドル）
try:
    import openai as openai_mod  # pip install openai
except Exception:
    openai_mod = None

# EasyOCR はオプション（未インストールでも落ちないように）
try:
    import easyocr  # pip install easyocr
except Exception:
    easyocr = None

# 画像系ライブラリは遅延・任意（未インストールを許容）
try:
    import numpy as np
    import cv2
    from PIL import Image
except Exception:
    np = None
    cv2 = None
    Image = None

# OpenAI v1
from openai import OpenAI

# ========== OpenAIクライアント生成（Azure互換） ==========
def get_openai_client(api_key: str) -> OpenAI:
    base = os.environ.get("OPENAI_BASE_URL")  # 例: https://{resource}.openai.azure.com/openai/v1
    if base:
        return OpenAI(api_key=api_key, base_url=base)
    return OpenAI(api_key=api_key)

# ========== ランタイム共通ストア（起動中のみ保持） ==========
@st.cache_resource(show_spinner=False)
def runtime_config():
    return {
        "common_password": None,   # 初回セットアップで管理者が設定
        "default_api_key": None,   # 任意：既定のAPIキー。未設定なら各ユーザーが毎回入力
    }

# ========== ログイン＆APIキー取得（毎回サイドバーで入力） ==========
def require_login_and_api() -> str:
    cfg = runtime_config()
    with st.sidebar:
        st.header("🔐 アクセス")

        # ⚙️ 管理者リセット（任意）
        with st.expander("⚙️ 管理者メニュー（リセット）"):
            reset_token = st.text_input("RESET と入力して有効化", key="reset_token")
            if st.button("初期セットアップをやり直す", key="btn_reset_setup"):
                if reset_token.strip().upper() != "RESET":
                    st.warning("RESET と入力してください。")
                else:
                    cfg["common_password"] = None
                    cfg["default_api_key"] = None
                    st.session_state.clear()
                    try:
                        st.rerun()
                    except Exception:
                        try:
                            st.experimental_rerun()
                        except Exception:
                            pass

        # 初回セットアップ：共通パスワードのみ設定
        if not cfg["common_password"]:
            st.info("初回セットアップ：共通パスワードのみ設定（APIキーは保存しません）")
            new_pw = st.text_input("共通パスワード（必須）", type="password", key="pw_setup")
            if st.button("保存", key="btn_save_pw"):
                if not new_pw:
                    st.error("共通パスワードは必須です。")
                else:
                    cfg["common_password"] = new_pw
                    cfg["default_api_key"] = None
                    st.success("セットアップ完了。以降はこのパスワードでログインできます。")
                    try:
                        st.rerun()
                    except Exception:
                        try:
                            st.experimental_rerun()
                        except Exception:
                            pass
            st.stop()

        # 通常ログイン：毎回 パスワード＋APIキー を入力
        pw = st.text_input("共通パスワードを入力", type="password", key="pw_login")
        user_key = st.text_input("OpenAI APIキー（必須）", type="password", key="user_api")

        if st.button("ログイン", key="btn_login"):
            if pw != cfg["common_password"]:
                st.error("パスワードが違います。")
                st.stop()
            if not user_key.strip():
                st.error("OpenAI APIキーは必須です。")
                st.stop()
            st.session_state["auth_ok"] = True
            st.session_state["user_api_key"] = user_key.strip()
            try:
                st.rerun()
            except Exception:
                try:
                    st.experimental_rerun()
                except Exception:
                    pass

        if not st.session_state.get("auth_ok"):
            st.stop()

    api_key = st.session_state.get("user_api_key", "")
    if not api_key:
        st.error("OpenAI APIキーが未入力です。サイドバーに入力してください。")
        st.stop()
    return api_key

# ========== 文字ユーティリティ ==========
def split_text_by_chars(text: str, chunk_size: int = 6000, overlap: int = 300) -> list[str]:
    text = text.strip()
    if len(text) <= chunk_size:
        return [text]
    chunks = []
    start = 0
    while start < len(text):
        end = min(len(text), start + chunk_size)
        cut = end
        for p in ("。", "！", "？", "\n"):
            idx = text.rfind(p, start, end)
            if idx != -1 and idx > start + 1000:
                cut = idx + 1
                break
        chunks.append(text[start:cut].strip())
        if cut >= len(text):
            break
        start = max(cut - overlap, 0)
    return [c for c in chunks if c]

def strip_timestamps(text: str) -> str:
    pattern = re.compile(
        r"^\s*\[\d{2}:\d{2}:\d{2}(?:\.\d{3})?\s*(?:→|->|-|－|—)\s*\d{2}:\d{2}:\d{2}(?:\.\d{3})?\]\s*",
        re.MULTILINE,
    )
    return pattern.sub("", text).strip()

# ========== FFmpeg/ffprobe 検出 ==========
PROJECT_DIR = Path(__file__).parent
FFBIN_CANDIDATES = [
    PROJECT_DIR / "ffmpeg-7.0.2-essentials_build" / "bin",
    Path(r"C:\\Users\\s-has\\Desktop\\動画音声原稿作成082025\\ffmpeg-7.0.2-essentials_build\\bin"),
    Path(r"C:\\Users\\s-has\\Desktop\\ffmpeg-7.0.2-essentials_build\\bin"),
]
FFMPEG_EXE = None
FFPROBE_EXE = None
for _bin in FFBIN_CANDIDATES:
    ff = _bin / "ffmpeg.exe"
    fp = _bin / "ffprobe.exe"
    if ff.exists():
        FFMPEG_EXE, FFPROBE_EXE = ff, (fp if fp.exists() else None)
        os.environ["PATH"] = str(_bin) + os.pathsep + os.environ.get("PATH", "")
        os.environ["FFMPEG_BINARY"] = str(ff)
        os.environ["IMAGEIO_FFMPEG_EXE"] = str(ff)
        AudioSegment.converter = str(ff)
        AudioSegment.ffmpeg = str(ff)
        if FFPROBE_EXE:
            AudioSegment.ffprobe = str(FFPROBE_EXE)
        break
else:
    ffmpeg_found = which("ffmpeg")
    ffprobe_found = which("ffprobe")
    if ffmpeg_found:
        FFMPEG_EXE = Path(ffmpeg_found)
        AudioSegment.converter = ffmpeg_found
        AudioSegment.ffmpeg = ffmpeg_found
    if ffprobe_found:
        FFPROBE_EXE = Path(ffprobe_found)
        AudioSegment.ffprobe = ffprobe_found

# ========== I/O ユーティリティ ==========
def save_uploaded_file_to_temp(uploaded_file) -> str:
    suffix = os.path.splitext(uploaded_file.name)[1]
    tmp_path = os.path.join(st.session_state["workdir"], f"upload_{int(time.time())}{suffix}")
    with open(tmp_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    return tmp_path

def ensure_wav(input_path: str) -> str:
    """アップロードされたファイルを 16kHz/mono の WAV に変換。
       pydub→失敗時は ffmpeg CLI にフォールバック。"""
    wav_path = os.path.splitext(input_path)[0] + "_16k.wav"
    # 1) まずは pydub
    try:
        audio = AudioSegment.from_file(input_path)
        audio = audio.set_channels(1).set_frame_rate(16000)
        audio.export(wav_path, format="wav")
        return wav_path
    except Exception:
        pass  # フォールバックへ

    # 2) フォールバック: ffmpeg CLI で “修復→音声抽出”
    ff = shutil.which("ffmpeg") or "ffmpeg"
    fixed_mp4 = os.path.splitext(input_path)[0] + "_fixed.mp4"
    try:
        p1 = subprocess.run(
            [ff, "-y", "-v", "error", "-i", input_path, "-c", "copy", "-movflags", "+faststart", fixed_mp4],
            stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, encoding="utf-8", errors="ignore"
        )
        src_for_audio = fixed_mp4 if os.path.exists(fixed_mp4) and p1.returncode == 0 else input_path
        p2 = subprocess.run(
            [ff, "-y", "-v", "error", "-i", src_for_audio,
             "-vn", "-ac", "1", "-ar", "16000", "-map", "0:a:0?",
             "-c:a", "pcm_s16le", wav_path],
            stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, encoding="utf-8", errors="ignore"
        )
        if p2.returncode == 0 and os.path.exists(wav_path):
            return wav_path
        raise RuntimeError(p2.stderr or "ffmpeg failed")
    except Exception as e2:
        st.error(
            "音声/動画の読み込みに失敗しました。\n"
            "（小さめのファイルで再アップロード、あるいは mp3/m4a でのアップをお試しください）\n\n"
            f"詳細: {e2}"
        )
        st.stop()

# ---- アップロード前の軽量化（まずはMP3 32kbps monoを試す）----
def shrink_audio_for_upload(src_path: str, target_mb: float = 50.0) -> tuple[str, float, str]:
    """
    return: (使うファイルパス, サイズMB, モード)
      モード: "original" / "mp3-32k" / "too_large"
    """
    try:
        size_mb = os.path.getsize(src_path) / (1024 * 1024)
    except Exception:
        return src_path, 0.0, "original"
    if size_mb <= target_mb:
        return src_path, size_mb, "original"
    # MP3 32kbps mono 再エンコード
    try:
        audio = AudioSegment.from_file(src_path)
        audio = audio.set_channels(1).set_frame_rate(16000)
        mp3_path = os.path.splitext(src_path)[0] + "_uploader.mp3"
        audio.export(mp3_path, format="mp3", bitrate="32k")
        size_mb2 = os.path.getsize(mp3_path) / (1024 * 1024)
        if size_mb2 <= target_mb:
            return mp3_path, size_mb2, "mp3-32k"
    except Exception:
        pass
    return src_path, size_mb, "too_large"

# ---- WAVを一定秒数で分割（デフォ 10分=600秒）----
def chunk_wav_by_time(src_wav: str, chunk_sec: int = 600) -> list[str]:
    audio = AudioSegment.from_file(src_wav)
    chunks = []
    base = os.path.splitext(src_wav)[0]
    total_ms = len(audio)
    step = chunk_sec * 1000
    i = 0
    for start_ms in range(0, total_ms, step):
        part = audio[start_ms:start_ms + step]
        out = f"{base}_part{i:03d}.wav"
        part.export(out, format="wav")
        chunks.append(out)
        i += 1
    return chunks

def format_timestamp(seconds: float) -> str:
    td = timedelta(seconds=float(seconds))
    total_seconds = int(td.total_seconds())
    ms = int((td.total_seconds() - total_seconds) * 1000)
    return f"{total_seconds//3600:02d}:{(total_seconds%3600)//60:02d}:{total_seconds%60:02d}.{ms:03d}"

def fmt_ts(x: float) -> str:
    return format_timestamp(x) if math.isfinite(x) else "…"

# ========== OpenAI で文字起こし ==========
def _safe_lang(forced_lang: str | None):
    if not forced_lang:
        return None
    lang = forced_lang.strip().lower()
    if lang in {"auto", "detect", "none", ""}:
        return None
    if len(lang) != 2:
        return None
    return lang

def transcribe_openai(wav_path: str, api_key: str, forced_lang: str | None = None):
    start = time.time()
    last_err = None
    try:
        if not api_key:
            raise ValueError("OpenAI API key is empty. Set OPENAI_API_KEY or provide api_key.")
        if not wav_path or not os.path.exists(wav_path):
            raise FileNotFoundError(f"Audio file not found: {wav_path}")

        file_size_mb = os.path.getsize(wav_path) / (1024 * 1024)
        # 50MB まで許可（>50MB は呼び出し側で分割してから渡すため通常到達しない）
        if file_size_mb > 50.0:
            raise RuntimeError(f"Audio file is too large ({file_size_mb:.1f} MB). Please split into smaller parts (<50MB).")

        client = get_openai_client(api_key)
        language = _safe_lang(forced_lang)

        with open(wav_path, "rb") as f:
            try:
                resp = client.audio.transcriptions.create(
                    model="gpt-4o-mini-transcribe",
                    file=f,
                    language=language
                )
            except Exception as e1:
                last_err = e1
                f.seek(0)
                resp = client.audio.transcriptions.create(
                    model="whisper-1",
                    file=f,
                    language=language
                )

        text = getattr(resp, "text", None) or (resp.get("text") if isinstance(resp, dict) else None)
        if not text:
            raise RuntimeError("OpenAI returned empty transcription text.")

        detected_lang = language or "auto"
        st.sidebar.info(f"Transcribed in {time.time()-start:.1f}s, size={file_size_mb:.1f}MB, lang={detected_lang}")
        segments = [{"start": 0.0, "end": 0.0, "text": text}]
        return segments, detected_lang

    except Exception as e:
        tb = traceback.format_exc()
        debug_blob = {
            "where": "transcribe_openai",
            "wav_path": wav_path,
            "forced_lang": forced_lang,
            "safe_lang": _safe_lang(forced_lang),
            "file_exists": os.path.exists(wav_path) if wav_path else False,
            "file_size_mb": (os.path.getsize(wav_path) / (1024*1024)) if (wav_path and os.path.exists(wav_path)) else None,
            "last_err_type": type(last_err).__name__ if last_err else None,
            "last_err_str": str(last_err) if last_err else None,
            "caught_err_type": type(e).__name__,
            "caught_err": str(e),
        }
        st.error("Transcription failed. See diagnostics below.")
        st.code(json.dumps(debug_blob, ensure_ascii=False, indent=2))
        st.code(tb)
        try:
            with open("/mount/src/transcribe_error.log", "a", encoding="utf-8") as logf:
                logf.write(json.dumps(debug_blob, ensure_ascii=False) + "\n")
                logf.write(tb + "\n")
        except Exception:
            pass
        raise RuntimeError(f"Transcription failed: {last_err or e}")

# ========== スライドと発話の対応付け ==========
def group_segments_by_slides(
    segments: List[Tuple[str, float, float]],
    slide_change_times: List[float]
) -> List[Dict[str, Any]]:
    last_end = max((e for _, _, e in segments), default=0.0)
    bounds = [0.0] + [t for t in slide_change_times if t < last_end] + [last_end]
    grouped = []
    for i in range(len(bounds)-1):
        start, end = bounds[i], bounds[i+1]
        bucket = []
        for t, s, e in segments:
            if e > start and s < end:
                bucket.append((t, max(s, start), min(e, end)))
        grouped.append({"index": i+1, "start": start, "end": end, "segments": bucket})
    return grouped

# ========== スライド抽出 & OCR ==========
def extract_slide_keyframes_with_times(video_path: str, out_dir: str, scene_thr: float=0.35) -> tuple[list[str], list[float]]:
    os.makedirs(out_dir, exist_ok=True)
    for p in glob.glob(os.path.join(out_dir, "*.jpg")):
        try:
            os.remove(p)
        except:
            pass

    ff_cmd = str(FFMPEG_EXE) if FFMPEG_EXE else (shutil.which("ffmpeg") or "ffmpeg")

    # 1) シーン変化抽出
    cmd = [
        ff_cmd, "-y", "-i", video_path,
        "-vf", f"select='gt(scene,{scene_thr})',showinfo",
        "-vsync", "vfr",
        os.path.join(out_dir, "%04d.jpg"),
    ]
    proc = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, encoding="utf-8", errors="ignore")
    stderr = proc.stderr or ""

    times = []
    for m in re.finditer(r"pts_time:([0-9]+\.[0-9]+)", stderr):
        try:
            times.append(float(m.group(1)))
        except:
            pass

    image_paths = sorted(glob.glob(os.path.join(out_dir, "*.jpg")))
    n = min(len(image_paths), len(times))
    if n > 0:
        return image_paths[:n], times[:n]

    # 2) フォールバック：3秒間隔で抽出
    for p in glob.glob(os.path.join(out_dir, "*.jpg")):
        try: os.remove(p)
        except: pass
    cmd_fb = [
        ff_cmd, "-y", "-i", video_path,
        "-vf", "fps=1/3,showinfo",
        "-vsync", "vfr",
        os.path.join(out_dir, "%04d.jpg"),
    ]
    subprocess.run(cmd_fb, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, encoding="utf-8", errors="ignore")

    image_paths = sorted(glob.glob(os.path.join(out_dir, "*.jpg")))
    if image_paths:
        approx_times = [i * 3.0 for i in range(len(image_paths))]
        return image_paths, approx_times

    # 3) それでも0なら先頭1枚だけ確保
    one_path = os.path.join(out_dir, "0001.jpg")
    subprocess.run([ff_cmd, "-y", "-ss", "00:00:01", "-i", video_path, "-vframes", "1", one_path],
                   stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, encoding="utf-8", errors="ignore")
    image_paths = sorted(glob.glob(os.path.join(out_dir, "*.jpg")))
    if image_paths:
        return image_paths, [0.0]

    return [], []

def _to_cv2_bgr(image_like):
    # 画像系ライブラリが無ければOCRはスキップ
    if (cv2 is None) or (np is None) or (Image is None):
        return None
    try:
        if isinstance(image_like, (bytes, bytearray)):
            arr = np.frombuffer(image_like, np.uint8)
            img = cv2.imdecode(arr, cv2.IMREAD_COLOR)
            return img
        if isinstance(image_like, str):
            img = cv2.imread(image_like, cv2.IMREAD_COLOR)
            if img is None:
                try:
                    pil = Image.open(image_like).convert("RGB")
                    return cv2.cvtColor(np.array(pil), cv2.COLOR_RGB2BGR)
                except Exception:
                    return None
            return img
        if isinstance(image_like, Image.Image):
            return cv2.cvtColor(np.array(image_like), cv2.COLOR_RGB2BGR)
        if isinstance(image_like, np.ndarray):
            if image_like.ndim == 2:
                return cv2.cvtColor(image_like, cv2.COLOR_GRAY2BGR)
            return image_like
    except Exception:
        return None
    return None

def _get_reader():
    """EasyOCR Reader を（あれば）キャッシュして使い回し。"""
    if easyocr is None:
        return None
    @st.cache_resource(show_spinner=False)
    def _cached_reader():
        return easyocr.Reader(['ja', 'en'], gpu=False)
    return _cached_reader()

def ocr_slides(image_paths: list) -> list[dict]:
    """
    image_paths: 画像パス/bytes/PIL/ndarray が混在していてもOK
    return: [{"index": i, "path": 元の参照, "text": 認識文字列}, ...]
    """
    if not image_paths:
        return []
    if easyocr is None or (cv2 is None) or (np is None) or (Image is None):
        # 依存がなければ空文字で返す（アプリは継続）
        return [{"index": i+1, "path": p, "text": ""} for i, p in enumerate(image_paths)]
    reader = _get_reader()
    results = []
    valid_found = False
    for idx, src in enumerate(image_paths, start=1):
        img = _to_cv2_bgr(src)
        if img is None or getattr(img, "size", 0) == 0:
            results.append({"index": idx, "path": src, "text": ""})
            continue
        valid_found = True
        try:
            lines = reader.readtext(img, detail=0)
            text = "\n".join(lines) if lines else ""
            results.append({"index": idx, "path": src, "text": text})
        except Exception:
            results.append({"index": idx, "path": src, "text": ""})
    if not valid_found:
        st.error("OCR用の画像を正しく読み込めませんでした（パス・形式・抽出処理をご確認ください）。")
    return results

# ========== 整形(生成AIなし) ==========
def to_verbatim_with_timestamps(segments: List[Tuple[str, float, float]]) -> str:
    lines: List[str] = []
    for t, s, e in segments:
        start_disp = format_timestamp(s) if math.isfinite(s) else "…"
        end_disp   = format_timestamp(e) if math.isfinite(e) else "…"
        lines.append(f"[{start_disp} → {end_disp}] {t}")
    return "\n".join(lines)

def heuristic_minutes(segments: List[Tuple[str, float, float]]) -> str:
    block, blocks, char_limit = [], [], 300
    for t, s, e in segments:
        if sum(len(x[0]) for x in block) + len(t) > char_limit and block:
            blocks.append(block); block = []
        block.append((t, s, e))
    if block: blocks.append(block)
    out = ["【議事録（自動整形・要点）】\n"]
    for i, b in enumerate(blocks, 1):
        out.append(f"■ トピック{i}（{format_timestamp(b[0][1])}–{format_timestamp(b[-1][2])}）")
        for t, _, _ in b: out.append(f"・{t}")
        out.append("")
    return "\n".join(out).strip()

def heuristic_abstract(segments: List[Tuple[str, float, float]]) -> str:
    text = " ".join(t for t, _, _ in segments)
    sentences = [s.strip() for s in text.replace("。", "。\n").splitlines() if s.strip()]
    return "【要旨（自動抽出）】\n" + "\n".join(sentences[:6])

def heuristic_article_academic(segments: List[Tuple[str, float, float]]) -> str:
    body = " ".join(t for t, _, _ in segments)
    lines = [
        "【学会報告記事（自動整形・AI不使用）】",
        "",
        "■ リード",
        "本講演では、演者が提示した主要ポイントを抜粋し、内容を簡潔に整理する。本文は自動整形のため、要点レベルの抜粋である。",
        "",
        "■ 背景・目的",
        "講演の背景、臨床上の意義、目的を本文から機械的に抽出・再構成。",
        "",
        "■ 方法・資料",
        "使用データ、対象、手法、評価指標などの記載を要点として抽出。",
        "",
        "■ 結果・所見",
        "本文から結果に相当する文を優先的に拾い上げ反映。",
        "",
        "■ 考察・結論",
        "臨床現場への示唆、限界、今後の展望を簡潔にまとめる。",
        "",
        "— 以下は逐語ベース本文（機械抽出） —",
        body,
    ]
    return "\n".join(lines)

def heuristic_guideline_commentary(slide_groups: List[Dict[str, Any]], ocr_notes: List[dict]) -> str:
    ocr_map = {o.get("index"): (o.get("text") or "").strip() for o in (ocr_notes or [])}
    lines = [
        "【ガイドライン解説（自動整形・AI不使用）】\n",
        "■ 背景",
        "・本解説は演者スライドとスピーチ内容を対応付けて再構成したもの。",
        "",
    ]
    for g in slide_groups:
        idx, ocr = g["index"], ocr_map.get(g["index"], "")
        lines.append(f"▼ Slide {idx}（{format_timestamp(g['start'])}–{fmt_ts(g['end'])}）")
        if ocr:
            title = ocr.splitlines()[0][:50]
            lines.append(f"【スライド要旨】{title}")
        for t, s, e in g["segments"][:6]:
            lines.append(f"・{t}")
        lines.append("")
    lines += ["■ 臨床への含意", "・本改訂により想定される診療上の影響点を要点化。", "", "■ 今後の課題", "・エビデンス強化が必要な論点、運用時の留意点。"]
    return "\n".join(lines).strip()

# ========== LLM（記事化/要旨/議事録） ==========
PURPOSE_PROMPTS = {
    "学会発表": (
        "以下の素材（音声逐語と任意のスライドOCR要約）から、学会報告記事を作成してください。"
        "見出し（導入/背景/目的/方法/結果/考察/結語）を付け、固有名詞と数値は改変せず、"
        "誇張や創作は避けてください。専門読者向けに簡潔で正確に。"
    ),
    "ガイドライン解説": (
        "以下の素材から、日本語のガイドライン改訂解説記事を作成してください。"
        "背景/改訂ポイント/推奨度・エビデンス/臨床への影響/課題/今後、の順に一度だけ骨組みを提示してください。"
        "テキストが複数パートに分割される場合でも、見出し・導入の再掲はしないでください。"
        "既出内容の再掲を避け、新規情報のみ追記する形で連続性を保ってください。"
        "英語は正確に日本語化し、引用は要旨化して書き直してください。"
    ),
    "ディスカッション": (
        "以下の素材から、ディスカッション記事を作成してください。"
        "論点整理/賛否の主張/根拠/一致点と相違点/結論と今後の検討課題、の順で、中立・簡潔にまとめてください。"
        "冗長な口語表現は削除し、方言は標準語に直してください。"
    ),
}

def llm_rewrite(kind: str, text: str, api_key: str | None,
                purpose: str | None = None,
                source_lang: str | None = None,
                target_lang: str | None = "ja") -> str:
    if openai_mod is None:
        return "[LLM未インストール] `pip install -U openai` を実行してください。"
    if not api_key:
        return "[APIキー未入力] サイドバーでAPIキーを入れてください。"

    sys_prompt = (
        "あなたは医学・医療系の日本語編集者です。臨床・学術文脈に沿って、"
        "読みやすく事実関係を保ったまま整文します。数値や引用は改変しません。"
    )
    pre = PURPOSE_PROMPTS.get(purpose or "学会発表", "")

    if (target_lang or "ja").lower() == "ja":
        lang_policy = (
            "最終出力は必ず日本語で書いてください。音声/スライドが日本語でない場合は正確に日本語へ翻訳し、"
            "専門用語は適切な日本語訳を用い、固有名詞・数値・単位は保持してください。"
        )
        if source_lang and str(source_lang).lower() != "ja":
            lang_policy += "（入力は日本語以外と検出されたため翻訳が必要です）"
    else:
        lang_policy = f"最終出力は必ず {target_lang} で書いてください。固有名詞・数値・単位は保持してください。"

    user_prompt_map = {
        "verbatim": "逐語記録（軽微な句読点整形のみ、意味改変禁止）：\n\n" + text,
        "minutes":  "議事録（見出し＋箇条書き、時系列）：\n\n" + text,
        "abstract": "学会抄録（目的/方法/結果/結論、600-900字）：\n\n" + text,
        "article":  "記事化（導入/背景/目的/方法/結果/考察/結語）：\n\n" + text,
    }
    if kind not in user_prompt_map:
        kind = "article"

    prompt = (pre + "\n\n" + lang_policy + "\n\n" + user_prompt_map[kind]).strip()

    client = openai_mod.OpenAI(api_key=api_key) if hasattr(openai_mod, "OpenAI") else None
    try:
        if client:
            resp = client.chat.completions.create(
                model="gpt-4o-mini-2024-07-18",
                messages=[{"role": "system", "content": sys_prompt},
                          {"role": "user", "content": prompt}],
                temperature=0.1,
            )
            result = resp.choices[0].message.content
        else:
            openai_mod.api_key = api_key
            resp = openai_mod.ChatCompletion.create(
                model="gpt-4o-mini-2024-07-18",
                messages=[{"role": "system", "content": sys_prompt},
                          {"role": "user", "content": prompt}],
                temperature=0.1,
            )
            result = resp["choices"][0]["message"]["content"]
    except Exception as e:
        return f"[LLMエラー] {e}"

    if kind != "verbatim":
        result = "【AI整形】\n" + result
    return result

def llm_translate_only(text: str, api_key: str | None,
                       source_lang: str | None = None,
                       target_lang: str = "ja") -> str:
    if openai_mod is None:
        return "[LLM未インストール] `pip install -U openai` を実行してください。"
    if not api_key:
        return "[APIキー未入力] サイドバーでAPIキーを入れてください。"

    sys_prompt = (
        "あなたは忠実な専門翻訳者です。以下のテキストを逐語的に日本語へ翻訳してください。"
        "要約・意訳・見出し付け・箇条書き化・体裁変更は行わないでください。"
        "段落や改行等の構造は可能な限り保持し、固有名詞・数値・単位は維持してください。"
    )
    if (target_lang or "ja").lower() != "ja":
        sys_prompt = sys_prompt.replace("日本語", target_lang)

    prompt = "【翻訳対象】\n" + text

    client = openai_mod.OpenAI(api_key=api_key) if hasattr(openai_mod, "OpenAI") else None
    try:
        if client:
            resp = client.chat.completions.create(
                model="gpt-4o-mini-2024-07-18",
                messages=[{"role": "system", "content": sys_prompt},
                          {"role": "user", "content": prompt}],
                temperature=0.0,
            )
            return resp.choices[0].message.content
        else:
            openai_mod.api_key = api_key
            resp = openai_mod.ChatCompletion.create(
                model="gpt-4o-mini-2024-07-18",
                messages=[{"role": "system", "content": sys_prompt},
                          {"role": "user", "content": prompt}],
                temperature=0.0,
            )
            return resp["choices"][0]["message"]["content"]
    except Exception as e:
        return f"[LLMエラー] {e}"

def llm_article_from_literal(literal_ja: str,
                             api_key: str | None,
                             purpose: str | None = "学会発表") -> str:
    if openai_mod is None:
        return "[LLM未インストール] `pip install -U openai` を実行してください。"
    if not api_key:
        return "[APIキー未入力] サイドバーでAPIキーを入れてください。"

    sys_prompt = (
        "あなたは医療・医学分野の編集者。入力は既に日本語へ逐語直訳された原稿。"
        "重複・言い換えの冗長だけを整理し、意味・事実は落とさず記事調（常体）に整える。"
        "【厳守】重複以外の削除禁止／数値・試験名・薬剤名・用量・単位は保持。"
        "見出しは『導入/背景/目的/方法/結果/考察/結語』の順で一度だけ。"
        "脚色・新情報の追加は禁止。"
        "文末は常体（〜だ／〜である）に統一し、です・ます調は禁止。"
    )
    preface = {
        "学会発表": "学会報告の速報トーンで、専門読者向けに簡潔で正確に。",
        "ガイドライン解説": "解説記事の文体で、背景→要点→臨床的含意を明確に。",
        "ディスカッション": "論点を明確化しつつ中立に記述。"
    }.get(purpose or "学会発表", "専門読者向けに簡潔で正確に。")

    user_prompt = (
        f"{preface}\n\n"
        "【入力（逐語直訳・日本語）】\n"
        + literal_ja.strip()
        + "\n\n【出力仕様】\n"
          "- TCROSS NEWS 学会発表記事のフォーマットに整形すること。\n"
          "- タイトルは「対象/疾患・介入: 試験名」とする。\n"
          "- 第1段落は「△△試験より、□□ことが、国、所属、演者名により、学会名とセッション名で発表された。」という形で書く（Conclusionの冒頭文を反映）。\n"
          "- 第2段落は試験デザインを記載（試験名、登録期間、国・施設数、患者数、群割付け、割付け数）。\n"
          "- 第3段落は患者背景を詳細に記載（差がなければ平均値で、年齢・性別・併存症・薬剤処方率を含める）。\n"
          "- 第4段落は主要評価項目の結果を記載（追跡期間、イベント率、HR、95%CI、p値を保持）。\n"
          "- 第5段落以降にサブ解析結果があれば記載。\n"
          "- 最終段落は演者のラストネームから始め、「…と、まとめた。」で必ず締める。\n"
          "- 同時掲載があれば「尚、△△試験は○○誌に掲載された。」と加える。\n"
          "- 記事調（常体）。\n"
          "- 見出しは『導入/背景/目的/方法/結果/考察/結語』。\n"
          "- 冗長な重複は統合。その他の内容は残す（削りすぎ禁止）。\n"
          "- 数値・用語はそのまま保持。\n"
          "- 箇条書きではなく段落ごとにまとめ、論理的な流れを持たせる。\n"
          "- 結果は逐語スクリプトの情報量を保持したまま記事調に整えること。\n"
    )

    client = openai_mod.OpenAI(api_key=api_key) if hasattr(openai_mod, "OpenAI") else None
    try:
        if client:
            resp = client.chat.completions.create(
                model="gpt-4o-mini-2024-07-18",
                messages=[{"role": "system", "content": sys_prompt},
                          {"role": "user", "content": user_prompt}],
                temperature=0.15,
            )
            return "【AI整形（直訳→記事調）】\n" + resp.choices[0].message.content
        else:
            openai_mod.api_key = api_key
            resp = openai_mod.ChatCompletion.create(
                model="gpt-4o-mini-2024-07-18",
                messages=[{"role": "system", "content": sys_prompt},
                          {"role": "user", "content": user_prompt}],
                temperature=0.15,
            )
            return "【AI整形（直訳→記事調）】\n" + resp["choices"][0]["message"]["content"]
    except Exception as e:
        return f"[LLMエラー] {e}"

# ========== DOCX 出力 ==========
def make_docx(title: str, content: str) -> bytes:
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Yu Gothic'
    font.size = Pt(11)

    doc.add_heading(title or "出力", level=1)
    for line in content.splitlines():
        if line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

# ========== Streamlit UI ==========
def main():
    st.set_page_config(page_title="InsighTCROSS® Smart Writer v11", layout="wide")

    # 1) ログイン＆APIキー入力（毎回）
    api_key = require_login_and_api()  # ← ここで毎回パスワード＆APIキーを入力
    st.session_state["api_key"] = api_key  # 接続テスト用に保持

    # 作業フォルダ
    if "workdir" not in st.session_state:
        st.session_state["workdir"] = os.path.abspath("./.work")
        os.makedirs(st.session_state["workdir"], exist_ok=True)

    # タイトル等
    st.title("InsighTCROSS® Smart Writer v11")
    if "transcript_text" not in st.session_state:
        st.session_state["transcript_text"] = ""
    if "generated_text" not in st.session_state:
        st.session_state["generated_text"] = ""
    st.write("音声/動画をアップロードして、逐語・直訳・議事録・要旨・記事に整形。動画はスライドOCR併用も可能。")

    # ===== サイドバー設定 =====
    with st.sidebar:
        st.header("設定")
        file_type = st.radio("ファイルタイプ", ["自動判定", "音声", "動画"], index=0, key="filetype")
        use_slide_ocr = st.toggle(
            "スライドOCRも併用（動画時）", value=False,
            help="スライドのキーフレームを抽出しOCRで文字も取り込みます（依存が無ければ空で継続）",
            key="toggle_ocr"
        )
        scene_sensitivity = st.slider("シーン変化感度", 0.10, 0.60, 0.35, 0.01, key="scene_thr")

        # 出力言語
        output_lang_label = st.selectbox("出力言語", ["日本語 (JPN)", "English (EN)"], index=0, key="out_lang")
        output_lang = "ja" if "JPN" in output_lang_label else "en"

        # 生成形式
        out_kind = st.selectbox(
            "出力タイプ",
            ["逐語(タイムスタンプ)", "直訳（日本語化のみ）", "議事録", "要旨", "記事", "ガイドライン解説"],
            key="out_kind"
        )
        purpose = st.selectbox("記事化の目的", ["学会発表", "ガイドライン解説", "ディスカッション"], index=0, key="purpose")
        attach_verbatim = st.toggle(
            "末尾に逐語原文を添付", value=False,
            help="原文言語の逐語テキストを末尾に付けます（通常はOFF推奨）",
            key="attach_verbatim"
        )

        # LLM整形のON/OFF（APIキーは require_login_and_api で受け取り済み）
        use_llm = st.toggle("生成AIで整形（任意）", value=False, key="use_llm")

        # 音声の言語（Whisperへの指示）
        speech_lang_label = st.selectbox("音声言語（Whisper）", ["英語", "日本語", "自動"], index=0, key="speech_lang")
        _lang_map = {"英語": "en", "日本語": "ja", "自動": None}
        forced_lang = _lang_map[speech_lang_label]

        # ---- 接続テスト ----
        st.divider()
        st.markdown("### 接続テスト")
        if st.button("🔎 OpenAI 接続テスト", key="btn_ping"):
            key = (st.session_state.get("api_key") or "").strip()
            if not key:
                st.error("先に APIキーを入力してください。")
            else:
                try:
                    c = get_openai_client(key)
                    _ = c.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[{"role": "user", "content": "ping"}],
                        max_tokens=5,
                        temperature=0.0,
                    )
                    st.success("OK: OpenAI へ到達できました。")
                except Exception as e:
                    st.error(
                        "NG: OpenAI へ接続/認証できません。ネットワーク or APIキーを確認してください。\n\n"
                        f"詳細: {e}"
                    )

        # ---- 10秒サンプルで転写テスト（任意）----
        st.markdown("### 転写ミニ診断")
        test_wav = "/mount/src/sample_10s.wav"
        if not os.path.exists(test_wav):
            st.caption(f"サンプル音声が見つかりません: {test_wav}（任意。置けばテストできます）")
        if st.button("🔎 10秒サンプルで転写テスト", key="btn_sample_transcribe"):
            if not os.path.exists(test_wav):
                st.error(f"サンプル音声が見つかりません: {test_wav}")
            else:
                try:
                    segs, lang = transcribe_openai(
                        test_wav,
                        api_key=api_key,
                        forced_lang="ja"
                    )
                    st.success(f"Sample OK. lang={lang}")
                    st.write(segs[0]["text"][:500])
                except Exception as e:
                    st.error(f"Sample failed: {e}")

    # ===== ファイルアップロード =====
    uploaded = st.file_uploader(
        "音声/動画ファイルをアップロード (mp3, m4a, wav, mp4, mov など)",
        type=["mp3","m4a","wav","mp4","mov","mkv","aac","flac"],
        key="uploader"
    )
    if not uploaded:
        return

    st.info(f"受信: {uploaded.name} / {uploaded.size/1024:.1f} KB")
    temp_path = save_uploaded_file_to_temp(uploaded)
    guessed = (uploaded.type or mimetypes.guess_type(uploaded.name)[0] or "")
    is_video = (st.session_state["filetype"] == "動画") or (st.session_state["filetype"] == "自動判定" and guessed.startswith("video/"))

    # 変換 → WAV 16kHz mono
    with st.spinner("変換中（WAV 16kHz mono）..."):
        wav_path = ensure_wav(temp_path)

    # 文字起こし：50MBまで直接。超過は自動で分割
    with st.spinner("🧠 OpenAIで文字起こし中…"):
        upload_path, up_mb, how = shrink_audio_for_upload(wav_path, target_mb=50.0)
        if how != "too_large":
            # そのまま（or mp3化）で1発転写
            segments, detected_lang = transcribe_openai(
                upload_path, api_key, forced_lang=forced_lang
            )
        else:
            # >50MB → 分割して連結
            st.warning(f"音声が {up_mb:.1f}MB と大きいため、10分刻みに分割してから転写します。")
            parts = chunk_wav_by_time(wav_path, chunk_sec=600)  # 10分
            segments, detected_lang = [], (forced_lang or "auto")
            offset = 0.0
            for i, p in enumerate(parts, start=1):
                st.caption(f"Part {i}/{len(parts)} を転写中…")
                segs_i, lang_i = transcribe_openai(p, api_key, forced_lang=forced_lang)
                try:
                    d_sec = AudioSegment.from_file(p).duration_seconds
                except Exception:
                    d_sec = 0.0
                txt = " ".join(s.get("text","") for s in segs_i)
                segments.append({"start": offset, "end": offset + d_sec, "text": txt})
                offset += d_sec

    st.success(f"文字起こし完了。セグメント数: {len(segments)} / 言語検出: {detected_lang}")

    # 逐語（タイムスタンプ付き）原稿（簡易：1セグメント合成前提）
    # 分割転写時は上でstart/endを詰めているので区切りが出ます
    def _to_triplets(segs_dicts):
        trips = []
        for s in segs_dicts:
            trips.append((s.get("text",""), float(s.get("start", 0.0)), float(s.get("end", 0.0))))
        return trips

    segments_triplets: List[Tuple[str, float, float]] = _to_triplets(segments)
    verbatim_text = to_verbatim_with_timestamps(segments_triplets)
    st.session_state["transcript_text"] = verbatim_text

    st.subheader("✍️ 逐語テキスト（編集可）")
    st.session_state["transcript_text"] = st.text_area(
        "逐語（必要に応じて修正してください）",
        value=st.session_state["transcript_text"],
        height=300,
        key="verbatim_editor"
    )

    # スライドOCR（任意）
    slide_groups, slide_notes, slide_digest = [], [], ""
    if is_video and st.session_state["toggle_ocr"]:
        with st.spinner("スライド抽出（キーフレーム+時刻）→ OCR 中..."):
            frames, slide_times = extract_slide_keyframes_with_times(
                video_path=temp_path,
                out_dir=os.path.join(st.session_state["workdir"], "slides"),
                scene_thr=st.session_state["scene_thr"],
            )

            st.write(f"抽出フレーム枚数: {len(frames)} / 切替検出: {len(slide_times)}")
            if frames:
                st.write("先頭3枚のパス:", frames[:3])
                try:
                    st.image(frames[0], caption="スライド抽出プレビュー（先頭）", use_container_width=True)
                except Exception as e:
                    st.warning(f"プレビュー表示に失敗: {e}")
            else:
                st.warning("抽出された画像が0枚です。フォールバックが効いていない可能性があります。")

            slide_notes = ocr_slides(frames)
            slide_groups = group_segments_by_slides(segments_triplets, slide_times)
            slide_digest = "\n\n".join(
                [f"[Slide {s['index']}]\n{s.get('text','')}" for s in slide_notes if s.get('text','').strip()]
            )
        st.success(f"スライド抽出: {len(slide_notes)} 枚 / 切替: {len(slide_times)} 点")

    edited_transcript = st.session_state["transcript_text"]
    cleaned_for_llm = strip_timestamps(edited_transcript)

    if st.session_state["out_kind"] == "ガイドライン解説" and slide_groups:
        chunks = []
        for g in slide_groups:
            idx = g["index"]
            ocr_text = ""
            for s in (slide_notes or []):
                if s.get("index") == idx:
                    ocr_text = (s.get("text") or "").strip()
                    break
            speech_text = "\n".join([t for (t, _, _) in g["segments"]])
            chunks.append(
                f"[Slide {idx} {format_timestamp(g['start'])}–{fmt_ts(g['end'])}]\n"
                f"<OCR>\n{ocr_text}\n</OCR>\n<SPEECH>\n{speech_text}\n</SPEECH>"
            )
        llm_source = "【スライド別素材】\n" + "\n\n".join(chunks)
    else:
        llm_source = cleaned_for_llm if not slide_digest else (
            f"【音声逐語】\n{cleaned_for_llm}\n\n【スライドOCR】\n{slide_digest}"
        )

    # 既定（ヒューリスティック）出力
    out_kind = st.session_state["out_kind"]
    if out_kind == "逐語(タイムスタンプ)":
        base_out = to_verbatim_with_timestamps(segments_triplets); kind_key = "verbatim"
    elif out_kind == "議事録":
        base_out = heuristic_minutes(segments_triplets); kind_key = "minutes"
    elif out_kind == "要旨":
        base_out = heuristic_abstract(segments_triplets); kind_key = "abstract"
    elif out_kind == "ガイドライン解説":
        base_out = heuristic_guideline_commentary(slide_groups, slide_notes) if slide_groups else \
                   "【ガイドライン解説（簡易）】\n" + heuristic_article_academic(segments_triplets)
        kind_key = "article"
    else:
        base_out = heuristic_article_academic(segments_triplets); kind_key = "article"

    final_out = base_out

    # ----- 生成AIで整形 -----
    st.markdown("---")
    st.subheader("🧠 生成AIで整形する")
    label_lang = "日本語" if output_lang == "ja" else "English"

    auto_generate = st.session_state["use_llm"]
    clicked = st.button(f"✨ 生成AIで整形（{label_lang}で出力）", key="btn_gen")
    do_generate = auto_generate or clicked

    if not do_generate:
        st.text_area("結果テキスト", value=final_out or "", height=400, key="no_gen_area")
        # ダウンロードだけは提供
        st.download_button("TXTダウンロード", data=final_out.encode("utf-8"), file_name="output.txt", key="dl_txt_nogen")
        docx_bytes = make_docx(title=f"{out_kind}（{purpose}）", content=final_out)
        st.download_button("DOCXダウンロード", data=docx_bytes, file_name="output.docx", key="dl_docx_nogen")
        return

    # 押下後
    if not st.session_state["use_llm"]:
        st.info("生成AIがOFFのため、ヒューリスティック整形の結果を表示します。")
        st.text_area("結果テキスト", value=final_out or "", height=400, key="gen_off_area")
        st.download_button("TXTダウンロード", data=final_out.encode("utf-8"), file_name="output.txt", key="dl_txt_off")
        docx_bytes = make_docx(title=f"{out_kind}（{purpose}）", content=final_out)
        st.download_button("DOCXダウンロード", data=docx_bytes, file_name="output.docx", key="dl_docx_off")
        return

    if not api_key:
        st.error("生成AIの整形には OpenAI APIキーが必要です（サイドバーで入力）。")
        st.stop()

    st.session_state.pop("ja_literal_for_article", None)

    final_out = base_out
    try:
        if out_kind == "逐語(タイムスタンプ)":
            with st.spinner("生成AIで整形中..."):
                final_out = llm_rewrite(
                    kind="verbatim",
                    text="【出力は必ず日本語】\n" + st.session_state["transcript_text"],
                    api_key=api_key,
                    purpose=purpose,
                    source_lang=detected_lang,
                    target_lang=output_lang,
                )
        elif out_kind == "直訳（日本語化のみ）":
            with st.spinner("直訳中..."):
                final_out = llm_translate_only(
                    text=cleaned_for_llm,
                    api_key=api_key,
                    source_lang=detected_lang,
                    target_lang="ja",
                )
        else:
            if out_kind == "記事" and (output_lang == "ja"):
                with st.spinner("英語→日本語 直訳 → 記事調 へ整形中..."):
                    ja_literal_for_article = llm_translate_only(
                        text=cleaned_for_llm,
                        api_key=api_key,
                        source_lang=detected_lang,
                        target_lang="ja",
                    )
                    final_out = llm_article_from_literal(
                        literal_ja=ja_literal_for_article,
                        api_key=api_key,
                        purpose=purpose,
                    )
                    st.caption("route: ARTICLE_FROM_LITERAL (ja) ✓ 直訳→記事調ルートを通過")
                    st.session_state["ja_literal_for_article"] = ja_literal_for_article
            else:
                llm_kind_call = {"議事録": "minutes", "要旨": "abstract"}.get(out_kind, "article")
                parts = split_text_by_chars(llm_source, chunk_size=6000, overlap=300)
                outs = []
                N = len(parts)
                for i, part in enumerate(parts, start=1):
                    meta = (
                        f"【分割パート {i}/{N}】\n"
                        "このパートでは新規情報のみを反映し、既出の見出しや導入は再掲しないでください。"
                    )
                    out_i = llm_rewrite(
                        kind=llm_kind_call,
                        text="【出力は必ず日本語】\n" + meta + "\n\n" + part,
                        api_key=api_key,
                        purpose=purpose,
                        source_lang=detected_lang,
                        target_lang=output_lang,
                    )
                    outs.append(out_i.strip())
                final_out = "\n\n".join(outs)
        st.success("生成AIでの整形が完了しました。")
    except Exception as e:
        st.error(f"整形に失敗しました: {e}")

    # ===== 三段表示 =====
    st.subheader("📝 原文（変更前・英語／タイムスタンプ除去）")
    st.text_area("原文", value=cleaned_for_llm, height=260, key="orig_area")

    st.subheader("🇯🇵 英語→日本語（直訳・整形なし）")
    if st.session_state["use_llm"] and api_key:
        cached_literal = st.session_state.get("ja_literal_for_article")
        if cached_literal:
            ja_literal = cached_literal
        else:
            with st.spinner("英語→日本語 直訳（プレビュー用）..."):
                ja_literal = llm_translate_only(
                    text=cleaned_for_llm,
                    api_key=api_key,
                    source_lang=detected_lang,
                    target_lang="ja",
                )
        st.text_area("直訳", value=ja_literal, height=260, key="literal_area")
    else:
        st.text_area("直訳", value="(生成AIがOFFまたはAPIキー未入力のため直訳は表示できません)", height=260, key="literal_off")

    st.subheader("📄 整形結果プレビュー")
    if out_kind == "ガイドライン解説" and output_lang == "ja" and final_out:
        for _p in ["背景", "改訂ポイント", "推奨度・エビデンス", "臨床への影響", "課題", "今後"]:
            final_out = re.sub(rf"(#+\s*{_p}\s*\n)(\s*\1)+", r"\1", final_out)
    st.text_area("整形結果", value=final_out, height=380, key="final_area")

    st.download_button("TXTダウンロード", data=final_out.encode("utf-8"), file_name="output.txt", key="dl_txt")
    docx_bytes = make_docx(title=f"{out_kind}（{purpose}）", content=final_out)
    st.download_button("DOCXダウンロード", data=docx_bytes, file_name="output.docx", key="dl_docx")

if __name__ == "__main__":
    main()

